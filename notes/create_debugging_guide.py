"""
Generate a beginner-friendly debugging guide .docx for the ArenaMCP project.
Explains every debugging step, jargon, and bash pipeline in plain English.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def set_style(doc):
    """Set up some basic styles to make the doc look clean."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text):
    """Add a grey-background code-style paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # We can't easily set background in python-docx, but the monospace + indent is enough.
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def create_guide():
    doc = Document()
    set_style(doc)

    # ==============================
    # TITLE PAGE
    # ==============================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ArenaMCP Debugging Guide')
    run.bold = True
    run.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('How We Got the MCP Server Working — Explained for Absolute Beginners')
    run.italic = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('May 2026')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ==============================
    # TABLE OF CONTENTS (manual)
    # ==============================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction — What This Guide Is',
        '2. What Is MCP? (The Big Picture)',
        '3. The Two Pipes — stdin and stdout',
        '4. Our Debugging Journey — Step by Step',
        '    4.1  Checking the Environment',
        '    4.2  Fixing the PYTHONPATH',
        '    4.3  Running the Server for the First Time',
        '    4.4  The Pipeline Test — Talking to the Server via Bash',
        '    4.5  Running the Automated Tests',
        '5. Jargon Buster — Dictionary of Every Term',
        '6. Command Reference — All Commands in One Place',
        '7. Troubleshooting Cheat Sheet',
    ]
    for item in toc_items:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ==============================
    # 1. INTRODUCTION
    # ==============================
    doc.add_heading('1. Introduction — What This Guide Is', level=1)
    doc.add_paragraph(
        'You asked us to build an MCP server that can tell you the top AI models on the Chatbot Arena leaderboard. '
        'We wrote the Python code, and then — like every real software project — we had to debug it before it actually worked. '
        'This guide walks through every single step we took, every command we ran, and every error we fixed.'
    )
    doc.add_paragraph(
        'More importantly, this guide explains every technical term as if you have never written a line of code. '
        'We do not assume you know what "stdin" means, what a "pipeline" is, or what "PYTHONPATH" does. '
        'We explain each one with real-world analogies.'
    )

    # ==============================
    # 2. WHAT IS MCP?
    # ==============================
    doc.add_heading('2. What Is MCP? (The Big Picture)', level=1)
    doc.add_paragraph(
        'MCP stands for "Model Context Protocol." That is a fancy name for a simple idea: '
        'it is a standard way for AI apps (like Claude Desktop) to ask a server for information.'
    )

    doc.add_heading('Imagine this scenario:', level=2)
    doc.add_paragraph(
        'You are at a restaurant. You (the AI app) want to know what soups are available. '
        'You do not go into the kitchen yourself. Instead, you ask a waiter (the MCP server). '
        'The waiter goes to the kitchen, gets the answer, and brings it back to you.'
    )
    add_bullet(doc, 'The AI app is the "client" — it asks questions.', 'Client: ')
    add_bullet(doc, 'The waiter is the "MCP server" — it receives questions and answers them.', 'Server: ')
    add_bullet(doc, 'The "menu" is the list of "tools" the server offers — things it knows how to do.', 'Tools: ')
    add_bullet(doc, 'The "order" is a "tool call" — a request like "get me the top 3 models."', 'Tool call: ')

    doc.add_paragraph(
        'In our project, the MCP server offers three tools: get_leaderboard, get_model_stats, and compare_models. '
        'When you ask "What are the top 3 AI models?", the AI app calls get_leaderboard(limit=3), '
        'the server looks up the data (currently from a mock table), and sends back a formatted answer.'
    )

    # ==============================
    # 3. STDIN AND STDOUT
    # ==============================
    doc.add_heading('3. The Two Pipes — stdin and stdout', level=1)
    doc.add_paragraph(
        'This is the most important concept to understand. MCP servers often communicate using something '
        'called "stdio transport." That sounds scary, but it is actually just "standard input" and "standard output."'
    )

    doc.add_heading('The Mailbox Analogy', level=2)
    doc.add_paragraph(
        'Imagine you have two mailboxes on your front door:'
    )
    add_bullet(doc, 'One is labeled INCOMING (stdin). The postman drops letters (data) into this box.', 'stdin (standard input): ')
    add_bullet(doc, 'The other is labeled OUTGOING (stdout). You put your replies into this box for the postman to collect.', 'stdout (standard output): ')

    doc.add_paragraph(
        'When an MCP server runs in "stdio mode," it reads requests from stdin (like reading a letter) '
        'and writes responses to stdout (like putting a reply in the outgoing box). '
        'The AI app on the other end does the opposite — it writes requests to the server\'s stdin '
        'and reads responses from the server\'s stdout.'
    )

    doc.add_heading('Why is this useful?', level=2)
    add_bullet(doc, 'No network needed — everything happens inside your computer, no internet required.')
    add_bullet(doc, 'Simple to test — you can talk to the server directly from the terminal using pipes (|).')
    add_bullet(doc, 'Fast — no HTTP overhead, just raw text back and forth.')

    doc.add_heading('What about stderr?', level=2)
    doc.add_paragraph(
        'There is a third mailbox called stderr (standard error). This is used ONLY for error messages and '
        'debugging logs. The server writes errors here so they do not get mixed up with the actual responses in stdout. '
        'Think of it as a "complaints box" — separate from the regular outgoing mail.'
    )

    doc.add_paragraph(
        'In our debugging, we used stderr to print status messages like '
        '"Server started successfully" while keeping stdout clean for actual JSON responses.'
    )

    # ==============================
    # 4. DEBUGGING JOURNEY
    # ==============================
    doc.add_heading('4. Our Debugging Journey — Step by Step', level=1)
    doc.add_paragraph(
        'Here is exactly what we did, in the order we did it, with every command we ran and every error we fixed.'
    )

    # 4.1
    doc.add_heading('4.1  Checking the Environment', level=2)
    doc.add_paragraph(
        'First, we needed to make sure Python and all the required packages were installed. '
        'Think of this like checking that the kitchen has all the ingredients before you start cooking.'
    )
    doc.add_paragraph('We ran:')
    add_code_block(doc, 'python3 --version')
    doc.add_paragraph('This checks which version of Python is installed. We need Python 3.10 or newer.')
    doc.add_paragraph('We also checked if the key package "mcp" was available:')
    add_code_block(doc, 'python3 -c "import mcp; print(mcp.__version__)"')
    doc.add_paragraph(
        'The "-c" flag means "run this piece of code." If Python can import mcp without errors, '
        'the package is installed. If it throws an error (ModuleNotFoundError), we need to install it.'
    )
    doc.add_paragraph('To install missing packages, we used:')
    add_code_block(doc, 'pip install -r requirements.txt')
    doc.add_paragraph(
        'This reads the file "requirements.txt" and installs every package listed there — '
        'like giving a shopping list to a personal shopper.'
    )

    # 4.2
    doc.add_heading('4.2  Fixing the PYTHONPATH', level=2)
    doc.add_paragraph(
        'This was one of the trickiest issues. PYTHONPATH is an environment variable — '
        'think of it as a list of folders that Python searches when you write "import something."'
    )
    doc.add_paragraph(
        'Imagine you tell a librarian "I want the book about ArenaClient." The librarian needs to know '
        'which section of the library to look in. PYTHONPATH is the list of sections the librarian searches.'
    )
    doc.add_paragraph(
        'Our server.py starts with:'
    )
    add_code_block(doc, 'from arena_mcp.arena_client import ArenaClient')
    doc.add_paragraph(
        'For Python to find the module "arena_mcp," the parent folder (src/) must be in PYTHONPATH. '
        'We set it like this:'
    )
    add_code_block(doc, 'export PYTHONPATH="/Users/mahedjaved/Downloads/AGENTIC/ArenaMCPProject/src"')
    doc.add_paragraph(
        'The "export" command sets an environment variable for the current terminal session. '
        'It is like telling the librarian "also look in this new section." Without this, '
        'Python throws "ModuleNotFoundError: No module named \'arena_mcp\'."'
    )
    doc.add_paragraph(
        'We put this in the run_arena_server.sh script so it is always set correctly when the server starts.'
    )

    # 4.3
    doc.add_heading('4.3  Running the Server for the First Time', level=2)
    doc.add_paragraph('We created a shell script to start the server:')
    add_code_block(doc, '#!/bin/bash\nexport PYTHONPATH="/Users/mahedjaved/Downloads/AGENTIC/ArenaMCPProject/src"\nexec /path/to/python -u /path/to/server.py')
    doc.add_paragraph(
        'Breaking this down:'
    )
    add_bullet(doc, '#!/bin/bash — This tells the computer "run this file using the bash shell." It is like saying "This recipe is meant for the oven, not the microwave."', '#! (shebang): ')
    add_bullet(doc, 'export PYTHONPATH=... — Sets the module search path (see section 4.2 above).', 'export: ')
    add_bullet(doc, 'exec python -u server.py — Starts the server with unbuffered output (-u flag, which means "send output immediately instead of waiting to collect it in batches").', 'exec: ')

    doc.add_paragraph(
        'When we first ran it, the server started but seemed to hang — no output appeared. '
        'That is actually correct! The server is waiting for input on stdin. '
        'It is like a waiter standing quietly, waiting for you to give your order.'
    )

    # 4.4
    doc.add_heading('4.4  The Pipeline Test — Talking to the Server via Bash', level=2)
    doc.add_paragraph(
        'To test whether the server was actually alive, we needed to send it a "JSON-RPC" request '
        'and see if it replied. JSON-RPC is just a structured way of asking a question:'
    )
    add_code_block(doc, '{"jsonrpc": "2.0", "method": "tools/list", "id": 1}')
    doc.add_paragraph(
        'This says: "Hello server, using JSON-RPC version 2.0, please list all your tools. My request ID is 1."'
    )
    doc.add_paragraph('We sent this to the server using a "pipeline":')
    add_code_block(doc, 'echo \'{"jsonrpc":"2.0","method":"tools/list","id":1}\' | PYTHONPATH=src python src/arena_mcp/server.py')
    doc.add_paragraph(
        'The "|" symbol is called a "pipe." Think of it as a physical tube:'
    )
    add_bullet(doc, 'On the left side of the pipe, echo prints our JSON message.', 'echo: ')
    add_bullet(doc, 'The pipe (|) sends that output directly into the server\'s stdin.', '| (pipe): ')
    add_bullet(doc, 'The server reads it, processes it, and writes a response to stdout.', 'server: ')
    add_bullet(doc, 'The response appears on our terminal screen.', 'stdout: ')

    doc.add_paragraph(
        'When this worked, we saw a JSON response listing three tools: get_leaderboard, get_model_stats, and compare_models. '
        'That confirmed the server was alive and the tools were registered correctly.'
    )

    doc.add_heading('What about errors? Capturing stderr', level=2)
    doc.add_paragraph(
        'Sometimes the server crashes or prints errors. Error messages go to stderr, not stdout. '
        'To see them, we redirect stderr to stdout:'
    )
    add_code_block(doc, 'echo \'...\' | PYTHONPATH=src python src/arena_mcp/server.py 2>&1')
    doc.add_paragraph(
        'The "2>&1" means "take stderr (file descriptor 2) and redirect it to stdout (file descriptor 1)." '
        'Think of it as: normally complaints (stderr) go into a separate bin. '
        'The "2>&1" says "actually, dump the complaints into the same bin as the regular output so I can see both."'
    )

    doc.add_paragraph('We also piped the output through a tool called Python\'s json.tool to make it readable:')
    add_code_block(doc, 'echo \'...\' | PYTHONPATH=src python src/arena_mcp/server.py | python3 -m json.tool')
    doc.add_paragraph(
        'This takes the raw JSON output and "pretty-prints" it — adding indentation and line breaks so a human can read it.'
    )

    # 4.5
    doc.add_heading('4.5  Running the Automated Tests', level=2)
    doc.add_paragraph(
        'Once the server was running, we ran the automated test suite to make sure everything worked correctly:'
    )
    add_code_block(doc, 'cd /Users/mahedjaved/Downloads/AGENTIC/ArenaMCPProject\nPYTHONPATH=src python -m pytest tests/ -v')

    doc.add_paragraph('Let us break down this command:')
    add_bullet(doc, 'PYTHONPATH=src — Tell Python where to find our modules (same as section 4.2).', 'PYTHONPATH=src: ')
    add_bullet(doc, 'python -m pytest — Run the "pytest" library as a script. pytest automatically finds and runs all test files.', 'python -m pytest: ')
    add_bullet(doc, 'tests/ — Look for tests in the tests/ folder.', 'tests/: ')
    add_bullet(doc, '-v — "verbose mode," meaning "show me details of every test that runs."', '-v: ')

    doc.add_paragraph('The tests check things like:')
    add_bullet(doc, 'Does get_leaderboard(3) return exactly 3 models?')
    add_bullet(doc, 'Does get_model_stats("gpt-4o") return rank 1 and Elo 1287?')
    add_bullet(doc, 'Does get_model_stats("ghost-model") say "not found"?')
    add_bullet(doc, 'Does compare_models() correctly compute which model has a higher Elo?')

    doc.add_paragraph('Each test shows either a green "PASSED" or red "FAILED." All four tests passed on our first run.')

    # ==============================
    # 5. JARGON BUSTER
    # ==============================
    doc.add_heading('5. Jargon Buster — Dictionary of Every Term', level=1)
    doc.add_paragraph('Here is every technical term used in this guide, explained in plain English.')

    terms = [
        ('Bash', 'A program (shell) that lets you type commands and run programs. Think of it as the "command center" of your computer — you type words, and it executes actions.'),
        ('CLI', 'Command-Line Interface. A way of interacting with the computer by typing commands instead of clicking buttons. The Terminal app is a CLI.'),
        ('echo', 'A command that prints text to the screen. Like shouting something — it just repeats what you give it.'),
        ('Environment variable', 'A named value stored in the system memory that programs can read. Like a post-it note on your computer saying "PYTHONPATH = /some/folder."'),
        ('exec', 'A Bash command that replaces the current shell process with a new program. Used in startup scripts to say "replace this script with the server process."'),
        ('File descriptor', 'A number that the operating system uses to refer to an open file or data stream. 0 = stdin, 1 = stdout, 2 = stderr.'),
        ('JSON', 'JavaScript Object Notation. A text format for structuring data, like a fill-in-the-blank form. Example: {"name": "gpt-4o", "elo": 1287}.'),
        ('JSON-RPC', 'A protocol for making requests using JSON. The message has a "method" (what to do) and "params" (arguments). Like calling a waiter and saying "I want the get_leaderboard tool with limit=3."'),
        ('Module', 'A single Python file containing code that can be reused. arena_client.py is a module.'),
        ('Package', 'A folder containing multiple Python modules, marked by an __init__.py file. arena_mcp/ is a package.'),
        ('Pipe (|)', 'A feature that connects the output of one command to the input of another. Like connecting a hose from one program\'s exhaust to another\'s air intake.'),
        ('pytest', 'A Python library that automatically finds and runs test functions, reporting which passed and which failed.'),
        ('PYTHONPATH', 'An environment variable listing folders where Python looks for modules. Without it, "import arena_mcp" fails because Python doesn\'t know where to find it.'),
        ('Redirect (2>&1)', 'A shell feature that sends output from one stream to another. "2>&1" means "send stderr (stream 2) to the same place as stdout (stream 1)."'),
        ('Shebang (#!)', 'The first line of a script (e.g., #!/bin/bash) that tells the operating system which interpreter to use. It is like a label on a recipe saying "this is for the oven."'),
        ('Shell', 'A program that interprets commands you type. Bash and Zsh are common shells.'),
        ('stdin (0)', 'Standard input — the data stream a program reads from. In our MCP server, stdin carries incoming JSON-RPC requests.'),
        ('stdout (1)', 'Standard output — the data stream a program writes its results to. In our MCP server, stdout carries outgoing JSON-RPC responses.'),
        ('stderr (2)', 'Standard error — a separate data stream for error messages and logs. Keeps errors from mixing with actual output.'),
        ('Transport', 'The mechanism by which two programs communicate. Our MCP server uses "stdio transport" (stdin/stdout). Other options include HTTP (web) and SSE (Server-Sent Events).'),
        ('Tool', 'A named function that an MCP server exposes. Our tools are get_leaderboard, get_model_stats, and compare_models.'),
        ('Unbuffered output (-u)', 'A Python flag that forces output to be sent immediately rather than waiting to fill a buffer (a temporary holding area). Like texting someone immediately instead of waiting to send 10 texts at once.'),
        ('Verbose (-v)', 'A flag that makes a program output more detail. Like asking a teacher to "show your work" instead of just giving the final answer.'),
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        run = p.add_run(term + ': ')
        run.bold = True
        p.add_run(definition)

    # ==============================
    # 6. COMMAND REFERENCE
    # ==============================
    doc.add_heading('6. Command Reference — All Commands in One Place', level=1)
    doc.add_paragraph('Here is every command we ran, with a plain-English explanation of what each one does.')

    commands = [
        ('Check Python version', 'python3 --version', 'Tells you which version of Python is installed.'),
        ('Check if a package is installed', 'python3 -c "import mcp; print(mcp.__version__)"', 'Tries to import the mcp library. If it works, prints the version. If not, throws an error.'),
        ('Install all dependencies', 'pip install -r requirements.txt', 'Reads requirements.txt and installs every package listed.'),
        ('Set PYTHONPATH', 'export PYTHONPATH="/path/to/src"', 'Tells Python where to find your modules. Must be run before starting the server.'),
        ('Start the MCP server (stdio mode)', 'python src/arena_mcp/server.py', 'Launches the server. It will wait quietly for input on stdin.'),
        ('List available tools (pipeline test)', 'echo \'{"jsonrpc":"2.0","method":"tools/list","id":1}\' | PYTHONPATH=src python src/arena_mcp/server.py', 'Sends a "list your tools" request to the server and prints the response.'),
        ('Call a tool with parameters', 'echo \'{"jsonrpc":"2.0","method":"tools/call","params":{"name":"get_leaderboard","arguments":{"limit":3}},"id":2}\' | PYTHONPATH=src python src/arena_mcp/server.py', 'Calls the get_leaderboard tool with limit=3.'),
        ('See errors too (2>&1)', 'echo \'...\' | PYTHONPATH=src python server.py 2>&1', 'Shows both normal output AND error messages.'),
        ('Pretty-print JSON output', 'echo \'...\' | PYTHONPATH=src python server.py | python3 -m json.tool', 'Makes the JSON response nicely indented and readable.'),
        ('Run all tests', 'PYTHONPATH=src python -m pytest tests/ -v', 'Runs all test files in the tests/ folder with detailed output.'),
        ('Run a single test file', 'PYTHONPATH=src python -m pytest tests/test_server.py -v', 'Runs only the server tests.'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'What It Does'
    hdr[1].text = 'Command'
    hdr[2].text = 'Plain English Explanation'
    for h in hdr:
        for p in h.paragraphs:
            for r in p.runs:
                r.bold = True

    for purpose, cmd, explanation in commands:
        row = table.add_row().cells
        row[0].text = purpose
        row[1].text = cmd
        row[2].text = explanation

    doc.add_paragraph()  # spacing

    # ==============================
    # 7. TROUBLESHOOTING CHEAT SHEET
    # ==============================
    doc.add_heading('7. Troubleshooting Cheat Sheet', level=1)
    doc.add_paragraph('If something goes wrong, here is how to diagnose and fix it.')

    problems = [
        ('"ModuleNotFoundError: No module named \'arena_mcp\'"', 
         'PYTHONPATH is not set or points to the wrong folder.',
         'Run: export PYTHONPATH="/Users/mahedjaved/Downloads/AGENTIC/ArenaMCPProject/src"'),
        ('"ModuleNotFoundError: No module named \'mcp\'"',
         'The MCP Python package is not installed.',
         'Run: pip install mcp'),
        ('Server starts but no output appears',
         'This is normal — it is waiting for stdin input.',
         'Send it a JSON-RPC request using the pipeline test (section 4.4).'),
        ('"Connection refused" or "Address already in use"',
         'A server is already running on that port (only matters for HTTP/SSE transport, not stdio).',
         'Kill the old process or use a different port. For stdio, this error should not occur.'),
        ('Tests fail with import errors',
         'PYTHONPATH was not set when running pytest.',
         'Run: PYTHONPATH=src python -m pytest tests/ -v'),
        ('JSON response says "Method not found"',
         'The tool name in the request does not match any registered tool.',
         'Check the exact tool name with the tools/list request first.'),
        ('Response is garbled or empty',
         'The server might have crashed before sending a response. Check stderr.',
         'Add 2>&1 to see error messages: ... 2>&1'),
        ('The server ran before but now does nothing',
         'A code change might have broken something.',
         'Run the tests to see what broke: PYTHONPATH=src python -m pytest tests/ -v'),
    ]

    for symptom, cause, fix in problems:
        p = doc.add_paragraph()
        run = p.add_run('Symptom: ')
        run.bold = True
        p.add_run(symptom)
        
        p = doc.add_paragraph()
        run = p.add_run('Likely cause: ')
        run.bold = True
        p.add_run(cause)
        
        p = doc.add_paragraph()
        run = p.add_run('Fix: ')
        run.bold = True
        p.add_run(fix)
        
        doc.add_paragraph()  # spacing between entries

    # ==============================
    # FOOTER
    # ==============================
    doc.add_paragraph('---')
    p = doc.add_paragraph()
    run = p.add_run('End of Guide')
    run.italic = True
    doc.add_paragraph(
        'Remember: every developer runs into these exact same problems. Debugging is not a sign of failure — '
        'it is the process of learning how your system actually works. Every error message is a clue, '
        'and every fix is a lesson.'
    )

    # ==============================
    # SAVE
    # ==============================
    output_path = '/Users/mahedjaved/Downloads/AGENTIC/ArenaMCPProject/notes/ArenaMCP_Debugging_Guide.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    create_guide()
